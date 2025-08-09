import os
import time
import re
import json
import numpy as np
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Tuple, Optional

import streamlit as st
import openai
try:
    import fitz  # PyMuPDF
    from PIL import Image
    from docx import Document
    import pytesseract
except ImportError as e:
    st.warning(f"Some optional dependencies missing: {e}. Install with: pip install pymupdf pillow python-docx pytesseract")

# -----------------------------
# Configuration & Utilities
# -----------------------------

# Add this near the top of your combined code, before the main app logic

def get_openai_api_key():
    """Get OpenAI API key from user input or existing session."""
    # Check if already in session
    if 'openai_api_key' in st.session_state and st.session_state.openai_api_key:
        return st.session_state.openai_api_key
    
    # Show API key input in sidebar
    with st.sidebar:
        st.markdown("## 🔑 OpenAI API Key")
        api_key = st.text_input(
            "Enter your OpenAI API key:",
            type="password",
            help="Get your API key from https://platform.openai.com/account/api-keys"
        )
        
        if api_key:
            st.session_state.openai_api_key = api_key
            st.success("✅ API key saved for this session")
            return api_key
        else:
            st.warning("Please enter your OpenAI API key to use the AI features")
            return None

# Then modify the ensure_openai_api_key function to use this:
def ensure_openai_api_key():
    """Set the OpenAI API key from user input."""
    api_key = get_openai_api_key()
    if api_key:
        openai.api_key = api_key
    else:
        st.error("OpenAI API key is required")
        st.stop()
# -----------------------------
# Document Processing
# -----------------------------

class DocumentProcessor:
    """Handles text extraction from various file formats."""
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(stream=BytesIO(file_bytes), filetype="pdf")
            return "\n".join([page.get_text() for page in doc])
        except Exception as e:
            return f"PDF extraction error: {e}"

    @staticmethod
    def extract_text_with_ocr(file_bytes: bytes) -> str:
        """Fallback OCR for PDFs with scanned images."""
        try:
            doc = fitz.open(stream=BytesIO(file_bytes), filetype="pdf")
            text = []
            for page in doc:
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text.append(pytesseract.image_to_string(img))
            return "\n".join(text)
        except Exception as e:
            return f"OCR error: {e}"

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception as e:
            return f"DOCX extraction error: {e}"

    @classmethod
    def extract_text(cls, file_bytes: bytes, filename: str) -> str:
        """Main extraction method that routes to appropriate parser."""
        if filename.lower().endswith(".pdf"):
            text = cls.extract_text_from_pdf(file_bytes)
            if len(text.strip()) < 50:  # Fallback to OCR if little text extracted
                text = cls.extract_text_with_ocr(file_bytes)
            return text
        elif filename.lower().endswith(".docx"):
            return cls.extract_text_from_docx(file_bytes)
        return ""

# -----------------------------
# Text Processing
# -----------------------------

class TextProcessor:
    """Handles text chunking, summarization, and analysis."""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for processing."""
        text = text.replace("\r\n", "\n")
        chunks = []
        start = 0
        N = len(text)
        while start < N:
            end = min(N, start + chunk_size)
            chunks.append(text[start:end].strip())
            start = end - overlap if (end - overlap) > start else end
        return [c for c in chunks if len(c) > 50]

    @staticmethod
    def summarize_text(text: str, max_length: int = 300, min_length: int = 100) -> str:
        """Generate a simple summary of the text."""
        if len(text.split()) < 50:
            return "Text too short to summarize."
        try:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
            summary = '. '.join(sentences[:4]) + '...'
            return summary if len(summary) > min_length else "Summary could not be generated."
        except Exception as e:
            return f"Summary generation failed: {e}"

    @staticmethod
    def analyze_contract_text(text: str, context_lines: int = 1) -> List[Dict]:
        """Analyze contract text for risky clauses using regex patterns."""
        findings = []
        seen_issues = set()

        # Pre-compiled patterns with metadata
        patterns = [
            {
                "regex": re.compile(r"\bliable for all damages\b|\bno limit on liability\b", re.IGNORECASE),
                "issue": "Unlimited Liability",
                "risk": "high",
                "suggestion": "Cap liability to the amount paid under the contract."
            },
            {
                "regex": re.compile(r"\bauto[- ]?renews\b|\bautomatically renews\b", re.IGNORECASE),
                "issue": "Automatic Renewal",
                "risk": "medium",
                "suggestion": "Add 30-day notice to opt-out before renewal."
            },
            {
                "regex": re.compile(r"\bmay terminate at any time\b|\bwithout cause\b", re.IGNORECASE),
                "issue": "One-Sided Termination",
                "risk": "medium",
                "suggestion": "Ensure both parties have equal termination rights."
            },
            {
                "regex": re.compile(r"\bgoverned by New York law\b|\bjurisdiction in London\b", re.IGNORECASE),
                "issue": "Foreign Jurisdiction",
                "risk": "high",
                "suggestion": "Use Ethiopian law and Addis Ababa courts for enforceability."
            },
        ]

        lines = text.splitlines()
        total_lines = len(lines)

        for i, line in enumerate(lines):
            line_clean = line.strip()
            for pattern in patterns:
                match = pattern["regex"].search(line_clean)
                if match:
                    issue_key = (i, pattern["issue"])
                    if issue_key in seen_issues:
                        continue  # Avoid duplicates
                    seen_issues.add(issue_key)

                    # Context lines before and after
                    start = max(0, i - context_lines)
                    end = min(total_lines, i + context_lines + 1)
                    context = "\n".join(lines[start:end]).strip()

                    # Highlight match in line
                    highlighted_line = pattern["regex"].sub(r"**\g<0>**", line_clean)

                    findings.append({
                        "line": i + 1,
                        "matched_text": highlighted_line,
                        "full_context": context,
                        "issue": pattern["issue"],
                        "risk": pattern["risk"],
                        "suggestion": pattern["suggestion"]
                    })
        return findings

# -----------------------------
# AI Services
# -----------------------------

class AIServices:
    """Handles all AI-related functionality including embeddings and LLM calls."""
    
    @staticmethod
    def get_embedding(texts: List[str], model: str = "text-embedding-3-small") -> np.ndarray:
        """Get embeddings for texts using OpenAI API."""
        resp = openai.Embedding.create(model=model, input=texts)
        return np.array([r["embedding"] for r in resp["data"]], dtype=np.float32)

    @staticmethod
    def cosine_similarity(a: np.ndarray, b: np.ndarray) -> np.ndarray:
        """Compute cosine similarity between embedding matrices."""
        a_norm = a / np.linalg.norm(a, axis=1, keepdims=True)
        b_norm = b / np.linalg.norm(b, axis=1, keepdims=True)
        return np.dot(a_norm, b_norm.T)

    @staticmethod
    def build_system_prompt() -> str:
        """Create the base system prompt for the LLM."""
        return (
            "You are LegalEase AI — a knowledgeable assistant about Ethiopian law. "
            "When answering, prefer Ethiopian laws where relevant, cite sources when possible, and when you use content "
            "from uploaded documents include the filename and a short quoted excerpt (<=200 chars). "
            "Be concise; structure answers as: Summary → Legal Context → Practical Advice → Sources → Disclaimer."
        )

    @staticmethod
    def build_user_prompt(user_question: str, retrieved: List[Tuple[str, str, float]], doc_context_first_n: int = 3) -> str:
        """Construct the user prompt with context."""
        prompt = f"User question: {user_question}\n\n"
        if retrieved:
            prompt += "Relevant document excerpts (most relevant first):\n"
            for fn, chunk, score in retrieved[:doc_context_first_n]:
                excerpt = chunk.replace("\n", " ")[:400]
                prompt += f"- From {fn} (score={score:.3f}): \"{excerpt}\"\n\n"
        prompt += "Answer following the structure: Summary; Legal Context (mention Ethiopian law/articles if known); Practical Advice; Sources; Disclaimer.\n"
        return prompt

    @classmethod
    def stream_chat_completion(cls, system_prompt: str, user_prompt: str, model: str = "gpt-4", temperature: float = 0.1, max_tokens: int = 800):
        """Stream chat completion from OpenAI with proper error handling."""
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        try:
            resp = openai.ChatCompletion.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
                stream=True
            )
        except Exception as e:
            st.error(f"LLM request failed: {e}")
            return

        partial = ""
        for chunk in resp:
            if "choices" in chunk:
                delta = chunk["choices"][0].get("delta", {})
                cont = delta.get("content", "")
                if cont:
                    partial += cont
                    yield cont

    @classmethod
    def generate_response(cls, prompt: str, context: Optional[str] = None, mode: str = "General Legal", tone: str = "Balanced") -> str:
        """Generate a response using either the local knowledge base or LLM."""
        # Ethiopian legal knowledge base
        legal_db = {
            "contract": {
                "validity": {
                    "source": "Ethiopian Civil Code Articles 1678-1690",
                    "content": "Four essential elements:\n1. Consent of the parties\n2. Capacity to contract\n3. Lawful object\n4. Lawful cause",
                    "examples": "Common issues: unclear terms, unequal bargaining power, illegal purposes"
                },
                "termination": {
                    "source": "Civil Code Articles 1806-1820 & Labor Proclamation No. 1156/2019",
                    "content": "Termination conditions vary by contract type:\n- Employment: Notice periods apply\n- Commercial: May include breach clauses",
                    "notice": "Reasonable notice typically required (30 days common)"
                }
            },
            "property": {
                "rights": {
                    "source": "Constitution Article 40 & Civil Code",
                    "content": "Right to private property guaranteed\nExpropriation requires fair compensation"
                }
            }
        }

        # Generate base response
        response = ""
        sources = []
        
        # Contract-related questions
        if "contract" in prompt.lower() or "agreement" in prompt.lower():
            if "valid" in prompt.lower() or "enforce" in prompt.lower():
                contract = legal_db["contract"]["validity"]
                response = f"Contract Validity Requirements:\n{contract['content']}\n\nExamples: {contract['examples']}"
                sources.append(contract["source"])
            elif "terminat" in prompt.lower() or "end" in prompt.lower():
                termination = legal_db["contract"]["termination"]
                response = f"Contract Termination Rules:\n{termination['content']}\n\nNotice: {termination['notice']}"
                sources.append(termination["source"])
        
        # Property-related questions
        elif "property" in prompt.lower() or "land" in prompt.lower():
            property = legal_db["property"]["rights"]
            response = f"Property Rights:\n{property['content']}"
            sources.append(property["source"])
        
        # Default response
        else:
            response = f"Regarding {prompt}, Ethiopian law generally requires... [provide general principles]"

        # Add document context if available
        if context:
            summary = TextProcessor.summarize_text(context)
            response = f"Document Context:\n- Summary: {summary[:200]}...\n\n{response}"
            
            findings = TextProcessor.analyze_contract_text(context)
            if findings:
                response += "\n\nKey Document Findings:"
                for f in findings[:3]:
                    response += f"\n- {f['issue']} (Risk: {f['risk'].upper()})"

        # Add legal sources
        if sources:
            response += "\n\nLegal Sources:\n- " + "\n- ".join(sources)

        # Apply tone
        if tone == "Simple":
            response = response.replace("typically", "usually")
            response = response.replace("proportionality", "fair share")
            response = response.replace("shall be considered", "should be thought of as")
        elif tone == "Formal":
            response = response.replace("can", "may")
            response = response.replace("should", "it is advisable to")
            response = "In accordance with legal principles, " + response

        # Final disclaimer
        response += "\n\nNote: This is general information, not legal advice. Consult a qualified attorney."

        return response

# -----------------------------
# Report Generation
# -----------------------------

class ReportGenerator:
    """Handles generation of PDF reports."""
    
    @staticmethod
    def clean_text_for_pdf(text: str) -> str:
        """Clean text to be ASCII-only for PDF generation."""
        if not text:
            return ""
        # First pass - replace common problematic characters
        replacements = {
            "—": "--", "–": "-", "’": "'", "“": '"', "”": '"',
            "‘": "'", "•": "*", "…": "...", "é": "e", "®": "(R)",
            "©": "(C)", "™": "(TM)", "°": " degrees "
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        # Second pass - remove any remaining non-ASCII
        return text.encode('ascii', 'ignore').decode('ascii')

    @staticmethod
    def generate_report_pdf(summary: str, findings: list) -> BytesIO:
        """Generate a PDF report with contract analysis results."""
        from fpdf import FPDF
        import io

        class PDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 16)
                self.cell(0, 10, ReportGenerator.clean_text_for_pdf('LegalEase AI - Contract Analysis Report'), ln=True, align='C')
                self.ln(10)

            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.cell(0, 10, ReportGenerator.clean_text_for_pdf(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M")} | Not Legal Advice'), align='C')

            def add_section(self, title, content):
                self.set_font('Helvetica', 'B', 14)
                self.cell(0, 10, ReportGenerator.clean_text_for_pdf(title), ln=True)
                self.set_font('Helvetica', '', 12)
                self.multi_cell(0, 6, ReportGenerator.clean_text_for_pdf(content))
                self.ln(8)

            def add_risk_table(self, findings):
                if not findings:
                    self.set_font('Helvetica', 'I', 12)
                    self.cell(0, 10, 'No risks detected.', ln=True)
                    self.ln(8)
                    return

                self.set_font('Helvetica', 'B', 14)
                self.cell(0, 10, 'Detected Risks', ln=True)
                self.set_font('Helvetica', 'B', 10)
                self.set_fill_color(220, 220, 220)
                self.cell(80, 8, 'Issue', 1, 0, 'C', 1)
                self.cell(30, 8, 'Risk', 1, 0, 'C', 1)
                self.cell(80, 8, 'Suggestion', 1, 1, 'C', 1)

                self.set_font('Helvetica', '', 9)
                for f in findings:
                    # Set highlight color based on risk level
                    fill_color = (255, 180, 180) if f["risk"] == "high" else (255, 240, 180)
                    self.set_fill_color(*fill_color)
                    
                    # Add cells with cleaned text
                    self.cell(80, 6, ReportGenerator.clean_text_for_pdf(f["issue"]), 1, 0, 'L', 1)
                    self.cell(30, 6, f["risk"].upper(), 1, 0, 'C', 1)
                    
                    # Clean and truncate suggestion if needed
                    suggestion = ReportGenerator.clean_text_for_pdf(f["suggestion"])
                    truncated_suggestion = (suggestion[:75] + "...") if len(suggestion) > 75 else suggestion
                    self.cell(80, 6, truncated_suggestion, 1, 1, 'L', 1)
                self.ln(10)

        # Create PDF in memory
        pdf = PDF()
        pdf.add_page()
        
        # Add content sections
        pdf.add_section("Document Summary", ReportGenerator.clean_text_for_pdf(summary))
        pdf.add_risk_table(findings)
        
        # Add disclaimer
        pdf.set_font('Helvetica', 'I', 10)
        pdf.set_text_color(150, 0, 0)
        pdf.multi_cell(0, 6, ReportGenerator.clean_text_for_pdf("DISCLAIMER: This report is AI-generated and not legal advice. Consult a licensed attorney."))

        # Create in-memory buffer
        buffer = io.BytesIO()
        buffer.write(pdf.output(dest='S').encode('latin-1'))
        buffer.seek(0)
        return buffer

# -----------------------------
# UI Components
# -----------------------------

class UIComponents:
    """Handles all UI rendering and layout."""
    
    @staticmethod
    def setup_page_config():
        """Configure the Streamlit page settings."""
        st.set_page_config(
            page_title="LegalEase AI",
            layout="wide",
            page_icon="⚖️"
        )

    @staticmethod
    def setup_sidebar():
        """Render the sidebar with settings and options."""
        with st.sidebar:
            st.title("⚙️ LegalEase AI Settings")
            
            # Theme selection
            theme = st.radio("🌓 Theme Mode", ["🌞 Light", "🌙 Dark"], key="theme_mode")
            
            # Model settings
            st.markdown("### 🤖 AI Configuration")
            model_choice = st.selectbox(
                "Model", 
                ["gpt-4", "gpt-4-turbo", "gpt-3.5-turbo"], 
                index=0,
                help="Select the AI model to use for responses"
            )
            emb_model = st.selectbox(
                "Embedding Model", 
                ["text-embedding-3-small", "text-embedding-3-large"], 
                index=0
            )
            temperature = st.slider(
                "Temperature", 
                0.0, 1.0, 0.1, 0.05,
                help="Controls randomness (0 = deterministic, 1 = creative)"
            )
            max_tokens = st.slider(
                "Max Response Length", 
                128, 2000, 800, step=64,
                help="Maximum number of tokens in the AI response"
            )
            
            # Chat settings
            st.markdown("### 💬 Chat Options")
            chatbot_mode = st.selectbox(
                "AI Mode",
                ["General Legal", "Contract Analysis", "Document Review"],
                help="Choose the AI assistant's focus"
            )
            chatbot_tone = st.select_slider(
                "Response Tone",
                options=["Formal", "Balanced", "Simple"],
                value="Balanced"
            )
            conversation_memory = st.checkbox(
                "Remember Conversation", 
                value=True,
                help="Maintain context across messages in this session"
            )
            enable_retrieval = st.checkbox(
                "Enable Document Retrieval", 
                value=True,
                help="Use uploaded documents for context when available"
            )
            
            # Resources
            st.markdown("---")
            st.markdown("### 📚 Legal Resources")
            st.markdown("- [Ethiopian Civil Code](https://chilot.me)")
            st.markdown("- [Contract Law Guidelines](https://chilot.me)")
            st.markdown("- [Legal Aid Services](https://ethiopianlaw.org)")
            
            return {
                "theme": theme,
                "model_choice": model_choice,
                "emb_model": emb_model,
                "temperature": temperature,
                "max_tokens": max_tokens,
                "chatbot_mode": chatbot_mode,
                "chatbot_tone": chatbot_tone,
                "conversation_memory": conversation_memory,
                "enable_retrieval": enable_retrieval
            }

    @staticmethod
    def apply_theme(theme: str):
        """Apply CSS styling based on selected theme."""
        if theme == "🌙 Dark":
            st.markdown("""
                <style>
                    body, .stApp {
                        background-color: #0e1117;
                        color: #fafafa;
                    }
                    .stButton>button, .stDownloadButton>button {
                        background-color: #1f2229;
                        color: white;
                        border-radius: 8px;
                        padding: 0.5em 1em;
                    }
                    .stTextInput>div>input {
                        background-color: #1f2229;
                        color: white;
                        border-radius: 5px;
                    }
                    .chat-message-user {
                        background-color: #1a3a4a !important;
                        padding: 1em;
                        border-radius: 10px;
                        color: #fafafa !important;
                    }
                    .chat-message-assistant {
                        background-color: #2a2a3a !important;
                        padding: 1em;
                        border-radius: 10px;
                        color: #fafafa !important;
                    }
                </style>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
                <style>
                    body, .stApp {
                        background-color: #ffffff;
                        color: #000000;
                    }
                    .stMarkdown, .stText, .stHeading, .stCaption, .stDataFrame {
                        color: #000000 !important;
                    }
                    .stButton>button, .stDownloadButton>button {
                        background-color: #f5f5f5;
                        color: #000000;
                        border-radius: 8px;
                        padding: 0.5em 1em;
                    }
                    .stTextInput>div>input {
                        background-color: #ffffff;
                        color: #000000;
                        border-radius: 5px;
                    }
                    .chat-message-user {
                        background-color: #e6f2ff !important;
                        padding: 1em;
                        border-radius: 10px;
                        color: #000000 !important;
                    }
                    .chat-message-assistant {
                        background-color: #f0f0f0 !important;
                        padding: 1em;
                        border-radius: 10px;
                        color: #000000 !important;
                    }
                </style>
            """, unsafe_allow_html=True)

    @staticmethod
    def setup_main_header():
        """Render the main header section."""
        st.markdown("<h1 style='font-size: 2.5em;'>💼 LegalEase AI</h1>", unsafe_allow_html=True)
        st.markdown("#### Empowering Ethiopian Legal Access Through AI")
        st.caption("⚠️ This tool does not provide legal advice. Always consult a licensed attorney.")
        st.markdown("---")

    @staticmethod
    def initialize_session_state():
        """Initialize all required session state variables."""
        if "doc_store" not in st.session_state:
            st.session_state.doc_store = []  # list of dicts: {"filename", "chunks": [...], "embeddings": np.array, "metadata": ...}

        if "chat_history" not in st.session_state:
            st.session_state.chat_history = [{
                "role": "assistant", 
                "content": "Hello! I'm LegalEase AI. I can help with:\n- Ethiopian contract law\n- Document review\n- Legal concepts\n\nHow can I assist you today?",
                "timestamp": datetime.now().isoformat()
            }]

        if "messages" not in st.session_state:
            st.session_state.messages = st.session_state.chat_history.copy()

    @staticmethod
    def render_chat_tab(settings: dict):
        """Render the chat interface tab."""
        with st.expander("💬 Smart Legal Assistant", expanded=True):
            st.markdown("Ask about Ethiopian civil law, contracts, rights, and more. You can also upload documents for analysis.")
            
            # Display chat messages
            for msg in st.session_state.messages:
                with st.chat_message(msg["role"]):
                    st.write(msg["content"])
                    if "timestamp" in msg:
                        st.caption(f"{datetime.fromisoformat(msg['timestamp']).strftime('%H:%M')}")
            
            # Document upload for chat context
            uploaded_file = st.file_uploader(
                "📎 Upload document for context (optional)",
                type=["pdf", "docx"],
                key="chat_file"
            )
            doc_context = ""
            findings = []
            if uploaded_file:
                with st.spinner("Processing document..."):
                    file_bytes = uploaded_file.read()
                    doc_context = DocumentProcessor.extract_text(file_bytes, uploaded_file.name)
                    if len(doc_context.strip()) > 50:
                        st.success(f"✅ Document loaded: {uploaded_file.name}")
                        findings = TextProcessor.analyze_contract_text(doc_context)
                    else:
                        st.warning("⚠️ Could not extract enough text from the document.")
            
            # Chat input
            if prompt := st.chat_input("Ask a legal question..."):
                # Add user message
                user_msg = {
                    "role": "user", 
                    "content": prompt,
                    "timestamp": datetime.now().isoformat()
                }
                st.session_state.messages.append(user_msg)
                st.session_state.chat_history.append(user_msg)
                
                with st.chat_message("user"):
                    st.write(prompt)
                    st.caption(f"{datetime.fromisoformat(user_msg['timestamp']).strftime('%H:%M')}")
                
                # Generate and display response
                with st.spinner("Analyzing..."):
                    message_placeholder = st.empty()
                    full_response = ""
                    
                    if settings["enable_retrieval"] and st.session_state.doc_store:
                        # Use RAG with document retrieval
                        retrieved = []
                        try:
                            q_emb = AIServices.get_embedding([prompt], model=settings["emb_model"])
                            for doc in st.session_state.doc_store:
                                sims = AIServices.cosine_similarity(q_emb, doc["embeddings"])[0]
                                top_idx = np.argmax(sims)
                                retrieved.append((doc["filename"], doc["chunks"][top_idx], float(sims[top_idx])))
                            retrieved = sorted(retrieved, key=lambda x: x[2], reverse=True)[:3]
                        except Exception as e:
                            st.error(f"Retrieval failed: {e}")
                            retrieved = []
                        
                        system_prompt = AIServices.build_system_prompt()
                        user_prompt = AIServices.build_user_prompt(prompt, retrieved)
                        
                        # Stream the response
                        for chunk in AIServices.stream_chat_completion(
                            system_prompt,
                            user_prompt,
                            model=settings["model_choice"],
                            temperature=settings["temperature"],
                            max_tokens=settings["max_tokens"]
                        ):
                            full_response += chunk
                            message_placeholder.markdown(full_response + "▌")
                        message_placeholder.markdown(full_response)
                    else:
                        # Use local knowledge base
                        response = AIServices.generate_response(
                            prompt,
                            context=doc_context if uploaded_file else None,
                            mode=settings["chatbot_mode"],
                            tone=settings["chatbot_tone"]
                        )
                        
                        # Simulate typing
                        for chunk in response.split('\n'):
                            full_response += chunk + '\n'
                            time.sleep(0.1)
                            message_placeholder.markdown(full_response + "▌")
                        message_placeholder.markdown(full_response)
                    
                    # Add assistant response
                    assistant_msg = {
                        "role": "assistant", 
                        "content": full_response,
                        "timestamp": datetime.now().isoformat()
                    }
                    st.session_state.messages.append(assistant_msg)
                    st.session_state.chat_history.append(assistant_msg)
            
            # Chat management buttons
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("🧹 Clear Chat", use_container_width=True):
                    st.session_state.messages = [{
                        "role": "assistant", 
                        "content": "Hello! I'm LegalEase AI. How can I help you with Ethiopian law today?",
                        "timestamp": datetime.now().isoformat()
                    }]
                    st.session_state.chat_history = st.session_state.messages.copy()
                    st.rerun()
            with col2:
                if st.session_state.messages:
                    chat_text = "\n".join([f"{m['role'].capitalize()}: {m['content']}" for m in st.session_state.messages])
                    st.download_button(
                        "📥 Save Chat", 
                        data=chat_text, 
                        file_name="legalease_chat.txt",
                        use_container_width=True
                    )
            with col3:
                example_questions = [
                    "What makes a contract valid in Ethiopia?",
                    "How can I terminate a lease agreement?",
                    "What are my property rights?"
                ]
                example = st.selectbox(
                    "💡 Example Questions",
                    ["Select an example..."] + example_questions,
                    index=0
                )
                if example and example != "Select an example...":
                    if "example_used" not in st.session_state or st.session_state.example_used != example:
                        st.session_state.example_used = example
                        st.session_state.messages.append({
                            "role": "user", 
                            "content": example,
                            "timestamp": datetime.now().isoformat()
                        })
                        st.session_state.chat_history = st.session_state.messages.copy()
                        st.rerun()

    @staticmethod
    def render_summarize_tab():
        """Render the document summarization tab."""
        with st.expander("📄 Summarize Legal Document", expanded=True):
            uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"], key="summarize_file")

            if uploaded_file:
                with st.spinner("🔍 Extracting text..."):
                    file_bytes = uploaded_file.read()
                    text = DocumentProcessor.extract_text(file_bytes, uploaded_file.name)

                if len(text.strip()) < 50:
                    st.warning("⚠️ Could not extract enough text. Is the file scanned or encrypted?")
                else:
                    st.success("✅ Text extracted!")
                    with st.expander("📄 Extracted Text (first 500 chars)"):
                        st.write(text[:500] + "...")

                    with st.spinner("🧠 Generating summary..."):
                        summary = TextProcessor.summarize_text(text)

                    st.markdown("### 📝 Summary")
                    st.write(summary)

    @staticmethod
    def render_risk_analyzer_tab():
        """Render the contract risk analysis tab."""
        with st.expander("🔍 Smart Contract Risk Analysis", expanded=True):
            uploaded_file = st.file_uploader("Upload PDF/DOCX", type=["pdf", "docx"], key="risk_file")

            if uploaded_file:
                with st.spinner("🔍 Scanning for risks..."):
                    file_bytes = uploaded_file.read()
                    text = DocumentProcessor.extract_text(file_bytes, uploaded_file.name)

                if len(text.strip()) < 50:
                    st.warning("⚠️ Not enough text extracted.")
                else:
                    findings = TextProcessor.analyze_contract_text(text)
                    summary = TextProcessor.summarize_text(text)

                    st.markdown("### 🎯 Filter Risks")
                    selected_levels = st.multiselect(
                        "Select risk levels to show", 
                        ["high", "medium"], 
                        default=["high", "medium"],
                        key="risk_levels"
                    )
                    filtered = [f for f in findings if f["risk"] in selected_levels]

                    if filtered:
                        st.metric("Total Issues", len(filtered))
                        for i, f in enumerate(filtered):
                            with st.expander(f"📌 {f['issue']} (Line {f['line']})"):
                                st.write(f"**Clause**: {f['matched_text']}")
                                st.write(f"**Risk Level**: {'🔴 High' if f['risk']=='high' else '🟡 Medium'}")
                                st.info(f"💡 **Suggestion**: {f['suggestion']}")
                    else:
                        st.success("✅ No major risks detected!")

                    try:
                        pdf_bytes = ReportGenerator.generate_report_pdf(summary, findings)
                        st.markdown("### 📄 Download Report")
                        st.download_button(
                            label="📥 Download Full Report (PDF)",
                            data=pdf_bytes,
                            file_name=f"legalease_analysis_{uploaded_file.name.split('.')[0]}.pdf",
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"Report generation failed: {e}")

    @staticmethod
    def render_clause_extractor_tab():
        """Render the clause extraction tab."""
        with st.expander("📚 Clause Extractor", expanded=True):
            uploaded_file = st.file_uploader("Upload Legal PDF/DOCX", type=["pdf", "docx"], key="clause_file")

            if uploaded_file:
                with st.spinner("Processing document..."):
                    file_bytes = uploaded_file.read()
                    text = DocumentProcessor.extract_text(file_bytes, uploaded_file.name)

                if len(text.strip()) < 50:
                    st.warning("⚠️ Not enough text extracted.")
                else:
                    st.success("✅ Extracted contract text.")
                    st.markdown("### 📑 Detected Clauses")

                    clauses = re.split(r"\n(?=\d+\.\s)", text)
                    for i, clause in enumerate(clauses):
                        if len(clause.strip()) > 20:
                            with st.expander(f"📄 Clause {i+1}"):
                                st.write(clause.strip())

    @staticmethod
    def render_document_upload_section():
        """Render the document upload section for RAG."""
        st.header("📄 Upload Document(s) for Context (optional)")
        uploaded_files = st.file_uploader(
            "Upload PDF or DOCX files (multiple allowed)", 
            type=["pdf", "docx"], 
            accept_multiple_files=True,
            key="rag_files"
        )

        if uploaded_files:
            with st.spinner("Processing files and creating embeddings..."):
                for f in uploaded_files:
                    file_bytes = f.read()
                    filename = f.name
                    text = DocumentProcessor.extract_text(file_bytes, filename)
                    if len(text.strip()) < 50:
                        st.warning(f"Could not extract text from {filename} or file too short.")
                        continue
                    
                    chunks = TextProcessor.chunk_text(text, chunk_size=1000, overlap=200)
                    if not chunks:
                        st.warning(f"No chunks for {filename}")
                        continue
                    
                    try:
                        embs = AIServices.get_embedding(chunks, model=st.session_state.get("emb_model", "text-embedding-3-small"))
                    except Exception as e:
                        st.error(f"Embedding error: {e}")
                        st.stop()
                    
                    st.session_state.doc_store.append({
                        "filename": filename,
                        "chunks": chunks,
                        "embeddings": embs,
                        "raw_text_preview": text[:2000]
                    })
                    st.success(f"Loaded & embedded {filename} ({len(chunks)} chunks).")

    @staticmethod
    def render_footer():
        """Render the application footer."""
        st.markdown("---")
        st.caption("Built by ASET (Alpha Software Engineering Tutorial) for Ethiopian Legal Empowerment • Capstone 2025")

# -----------------------------
# Main App
# -----------------------------

def main():
    """Main application entry point."""
    # Setup UI components
    UIComponents.setup_page_config()
    settings = UIComponents.setup_sidebar()
    UIComponents.apply_theme(settings["theme"])
    UIComponents.setup_main_header()
    UIComponents.initialize_session_state()
    
    # Create tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "🗨️ Smart Chat",
        "📄 Summarize Document",
        "🔍 Analyze Legal Risks",
        "📚 Extract Contract Clauses"
    ])
    
    # Render tabs
    with tab1:
        UIComponents.render_chat_tab(settings)
    
    with tab2:
        UIComponents.render_summarize_tab()
    
    with tab3:
        UIComponents.render_risk_analyzer_tab()
    
    with tab4:
        UIComponents.render_clause_extractor_tab()
    
    # Additional sections
    UIComponents.render_document_upload_section()
    UIComponents.render_footer()

if __name__ == "__main__":
    main()
